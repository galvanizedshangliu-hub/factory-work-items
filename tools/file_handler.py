"""
钉钉文件消息处理 — 下载 + Excel/CSV 解析
用于生产计划单的输入文件处理
"""
import os, json, io, requests, platform
from datetime import datetime

DINGTALK_OPENAPI = "https://api.dingtalk.com"


def log(msg):
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] [file_handler] {msg}", flush=True)


def get_access_token(client_id: str, client_secret: str) -> str:
    """获取钉钉企业内部应用 access_token"""
    resp = requests.post(
        f"{DINGTALK_OPENAPI}/v1.0/oauth2/accessToken",
        headers={"Content-Type": "application/json"},
        json={"appKey": client_id, "appSecret": client_secret},
        timeout=10,
    )
    resp.raise_for_status()
    data = resp.json()
    return data["accessToken"]


def get_file_download_url(download_code: str, client_id: str, access_token: str) -> str:
    """
    通过 downloadCode 换取文件下载链接（含权限，有效期有限）
    参考：https://open.dingtalk.com/document/isvapp/download-the-file-content-of-the-robot-receiving-message
    """
    headers = {
        "Content-Type": "application/json",
        "Accept": "*/*",
        "x-acs-dingtalk-access-token": access_token,
        "User-Agent": (
            f"DingTalkStream/1.0 SDK/0.1.0 Python/{platform.python_version()} "
            f"(+https://github.com/open-dingtalk/dingtalk-stream-sdk-python)"
        ),
    }
    body = {
        "robotCode": client_id,
        "downloadCode": download_code,
    }
    url = f"{DINGTALK_OPENAPI}/v1.0/robot/messageFiles/download"
    resp = requests.post(url, headers=headers, json=body, timeout=15)
    resp.raise_for_status()
    return resp.json()["downloadUrl"]


def download_file_content(download_url: str) -> bytes:
    """从下载链接获取文件内容"""
    resp = requests.get(download_url, timeout=30)
    resp.raise_for_status()
    return resp.content


def parse_spreadsheet(file_bytes: bytes, filename: str) -> dict:
    """
    解析 Excel (.xlsx/.xls) 或 CSV 文件
    返回 {"success": bool, "sheets": {sheet_name: [rows]}, "row_count": int, "columns": [str]}
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext in (".xlsx", ".xls"):
        return _parse_excel(file_bytes, filename)
    elif ext == ".csv":
        return _parse_csv(file_bytes, filename)
    elif ext == ".docx":
        return _parse_docx(file_bytes, filename)
    else:
        return {"success": False, "error": f"不支持的文件格式: {ext}"}


def _parse_excel(file_bytes: bytes, filename: str) -> dict:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    sheets = {}
    total_rows = 0
    all_columns = []

    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # 跳过全空行
            if all(v is None for v in row):
                continue
            rows.append([str(v) if v is not None else "" for v in row])

        sheets[name] = rows
        total_rows += len(rows)
        if rows and not all_columns:
            all_columns = rows[0]

    wb.close()
    return {
        "success": True,
        "sheets": sheets,
        "row_count": total_rows,
        "columns": all_columns,
        "sheet_names": list(sheets.keys()),
    }


def _parse_csv(file_bytes: bytes, filename: str) -> dict:
    import csv

    text = file_bytes.decode("utf-8-sig")  # 兼容 BOM
    reader = csv.reader(io.StringIO(text))
    rows = [row for row in reader if any(cell.strip() for cell in row)]

    return {
        "success": True,
        "sheets": {"Sheet1": rows},
        "row_count": len(rows),
        "columns": rows[0] if rows else [],
        "sheet_names": ["Sheet1"],
    }


def _parse_docx(file_bytes: bytes, filename: str) -> dict:
    """解析 Word (.docx) 文件中的表格和图片"""
    from docx import Document
    import zipfile

    doc = Document(io.BytesIO(file_bytes))
    sheets = {}
    total_rows = 0
    all_columns = []

    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # 跳过全空行
            if all(c == "" for c in cells):
                continue
            rows.append(cells)

        sheet_name = f"表格{ti + 1}"
        if rows:
            sheets[sheet_name] = rows
            total_rows += len(rows)
            if not all_columns:
                all_columns = rows[0]

    # 如果没有任何表格，尝试提取纯文本
    if not sheets:
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            sheets = {"正文": [[p] for p in paragraphs]}
            total_rows = len(paragraphs)
            all_columns = ["内容"]

    # ── 提取文档中的嵌入图片 ──
    images = []
    try:
        zf = zipfile.ZipFile(io.BytesIO(file_bytes))
        for name in zf.namelist():
            if name.startswith("word/media/") and any(
                name.lower().endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")
            ):
                images.append({
                    "name": os.path.basename(name),
                    "path": name,
                    "size": zf.getinfo(name).file_size,
                    "data": zf.read(name),
                })
        zf.close()
    except Exception:
        pass

    return {
        "success": True,
        "sheets": sheets,
        "row_count": total_rows,
        "columns": all_columns,
        "sheet_names": list(sheets.keys()),
        "images": images,  # 新增：文档中的图片列表
    }


def format_file_summary(parsed: dict, filename: str) -> str:
    """生成文件解析摘要，方便在群聊中预览"""
    if not parsed["success"]:
        return f"📎 {filename}\n❌ 解析失败: {parsed.get('error', '未知错误')}"

    ext = os.path.splitext(filename)[1].lower()
    region_label = "个区域" if ext == ".docx" else "个工作表"

    lines = [
        f"📎 {filename}",
        f"共 {parsed['row_count']} 行数据，{len(parsed['sheet_names'])}{region_label}",
    ]
    for sheet_name, rows in parsed["sheets"].items():
        n = len(rows)
        preview = rows[0] if rows else []
        cols = ", ".join(preview[:8])
        if len(preview) > 8:
            cols += f" ... 等{len(preview)}列"
        lines.append(f"  ▸ {sheet_name}: {n}行 | 列: {cols}")

    # 图片数量
    img_count = len(parsed.get("images", []))
    if img_count > 0:
        total_kb = sum(img.get("size", 0) for img in parsed["images"]) / 1024
        lines.append(f"  🖼️ 文档含 {img_count} 张图片（共 {total_kb:.0f}KB）")

    return "\n".join(lines)


def handle_file_message(extensions: dict, client_id: str, client_secret: str) -> dict:
    """
    处理文件消息的入口函数。
    从 ChatbotMessage 的 extensions 中提取文件信息，下载并解析。

    返回:
        {"success": bool, "filename": str, "parsed": dict, "summary": str, "error": str}
    """
    # 从 extensions 提取文件信息
    content = extensions.get("content", {})
    download_code = content.get("downloadCode", "")
    filename = content.get("fileName", "unknown")

    if not download_code:
        return {"success": False, "error": "消息中未找到 downloadCode", "filename": filename}

    log(f"📎 收到文件: {filename} (downloadCode: {download_code[:20]}...)")

    try:
        # 1. 获取 access_token
        token = get_access_token(client_id, client_secret)

        # 2. 换取下载链接
        download_url = get_file_download_url(download_code, client_id, token)
        log(f"🔗 获取下载链接成功")

        # 3. 下载文件内容
        file_bytes = download_file_content(download_url)
        log(f"📥 下载完成: {len(file_bytes)} bytes")

        # 4. 解析文件
        parsed = parse_spreadsheet(file_bytes, filename)
        summary = format_file_summary(parsed, filename)
        log(f"📊 解析完成: {parsed.get('row_count', 0)}行")

        # 5. 保存本地副本（调试用）
        try:
            safe_name = filename.replace("/", "_").replace("\\", "_")
            local_path = os.path.join(os.path.dirname(__file__), "..", "data", "downloaded_" + safe_name)
            with open(local_path, "wb") as f:
                f.write(file_bytes)
            log(f"💾 已保存本地副本: {local_path}")
        except Exception:
            pass

        return {
            "success": True,
            "filename": filename,
            "parsed": parsed,
            "summary": summary,
            "file_bytes": file_bytes,
        }

    except Exception as e:
        log(f"❌ 文件处理失败: {e}")
        return {"success": False, "error": str(e), "filename": filename}
